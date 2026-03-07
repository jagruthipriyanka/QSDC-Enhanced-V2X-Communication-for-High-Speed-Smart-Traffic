from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # --- Header Styling ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Autonomous-Flow: Technical Manual & User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1 ---
    doc.add_heading('1. Project Overview & Core Objective', level=1)
    p = doc.add_paragraph(
        'Autonomous-Flow is a next-generation simulation platform that merges '
        'Quantum Communication, V2X (Vehicle-to-Everything) Traffic Systems, '
        'and Machine Learning.'
    )
    p = doc.add_paragraph(
        'The core objective is to create an urban environment where autonomous vehicles (AVs) '
        'communicate using "eavesdrop-proof" quantum channels while using artificial intelligence '
        'to predict network failures and prevent collisions.'
    )

    # --- Section 2 ---
    doc.add_heading('2. Key Terminologies & Concepts', level=1)
    
    doc.add_heading('A. Quantum Superdense Coding (QSDC)', level=2)
    doc.add_paragraph('What it is: A protocol that allows two classical bits of information to be sent using only one qubit, provided the sender and receiver share an entangled pair.', style='List Bullet')
    doc.add_paragraph('Bell State (|Φ⁺⟩): The "entangled" state where two qubits are linked. What happens to one instantly affects the other.', style='List Bullet')
    doc.add_paragraph('Pauli Encoding: The method Alice (the vehicle) uses to change her qubit\'s state (X or Z gates) to "hide" the message inside.', style='List Bullet')

    doc.add_heading('B. Security & QBER', level=2)
    doc.add_paragraph('QBER (Quantum Bit Error Rate): The percentage of bits that were decoded incorrectly.', style='List Bullet')
    doc.add_paragraph('Threshold (11%): In quantum physics, if the error rate goes above 11%, it mathematically proves that someone (Eve) is trying to "listen" or measure the qubits mid-flight.', style='List Bullet')

    doc.add_heading('C. V2X & mmWave', level=2)
    doc.add_paragraph('V2X: Vehicle-to-Everything communication.', style='List Bullet')
    doc.add_paragraph('mmWave (5G/6G): High-frequency radio waves (28-60GHz) that provide massive speed but are very sensitive to distance and movement.', style='List Bullet')
    doc.add_paragraph('SNR (Signal-to-Noise Ratio): How "clean" the signal is. A low SNR means the connection is about to drop.', style='List Bullet')

    # --- Section 3 ---
    doc.add_heading('3. How It Works (The Engine)', level=1)
    doc.add_paragraph('Step 1: Mobility & Environment', style='List Number')
    doc.add_paragraph('The simulation moves cars through a 2D grid. Each car calculates its distance to the nearest "Roadside Unit" (Tower). As a car moves further away or travels faster, the "noise" in the quantum channel increases.', style='Normal')
    
    doc.add_paragraph('Step 2: Quantum Transmission', style='List Number')
    doc.add_paragraph('When a car sends a message, a Bell Pair is created, encoded via Pauli gates, transmitted through a noisy channel, and measured by the Tower to retrieve the bits.', style='Normal')
    
    doc.add_paragraph('Step 3: AI Prediction', style='List Number')
    doc.add_paragraph('While the cars move, a Random Forest model watches the SNR levels. If it sees a pattern that looks like an upcoming signal drop, it flags a warning in the dashboard.', style='Normal')

    # --- Section 4 ---
    doc.add_heading('4. Understanding the Dashboard (UI Components)', level=1)
    doc.add_paragraph('The Simulation Canvas: Renders AVs (triangles) and RSUs (towers). Vehicles glow neon purple during quantum transmission.', style='List Bullet')
    doc.add_paragraph('KPI Cards: Displays Q-Fidelity (aim for 1.0), Latency (ms), and Collision Risks.', style='List Bullet')
    doc.add_paragraph('Security Console: Logs every packet. Red text indicates a "BREACH DETECTED" (QBER > 11%).', style='List Bullet')
    doc.add_paragraph('ML Intelligence: Monitors the status of the SNR Predictor and Risk Detector models.', style='List Bullet')

    # --- Section 5 ---
    doc.add_heading('5. Step-by-Step Procedure', level=1)
    proc = [
        "Open the website (http://localhost:5000).",
        "Click the [Start] button to begin the live simulation.",
        "Observe the vehicles moving and the charts updating real-time.",
        "Use the QSDC Transmitter card to send a manual message.",
        "Toggle 'Simulate Eavesdropping' to see the security system detect Eve.",
        "Monitor the Security Console for AI-based connection drop predictions.",
        "Verify the historical data in your Supabase Table Editor."
    ]
    for step in proc:
        doc.add_paragraph(step, style='List Number')

    # --- Section 6 ---
    doc.add_heading('6. Expected Outcomes', level=1)
    doc.add_paragraph('Zero-Trust Security: Reliable detection of eavesdropping via QBER spikes.', style='List Bullet')
    doc.add_paragraph('Proactive Network Management: AI-driven warnings of upcoming signal failures.', style='List Bullet')
    doc.add_paragraph('Real-time Urban Visualization: High-fidelity monitoring of AV communication ecosystems.', style='List Bullet')

    # --- Save ---
    doc.save(r'd:\QC\Autonomous-Flow_Manual.docx')
    print("Document created successfully at d:\QC\Autonomous-Flow_Manual.docx")

if __name__ == "__main__":
    create_manual()
